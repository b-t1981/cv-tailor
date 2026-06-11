import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.config import settings
from app.services.docx_processor import save_docx
from app.services.pdf_exporter import export_docx_to_pdf


def export_cover_letter_docx(
    session_id: str,
    cover_letter: str,
    *,
    company_name: str = "",
    job_title: str = "",
) -> tuple[str, str]:
    """Build a DOCX cover letter and return (filename, download_url)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if company_name or job_title:
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("\n".join(part for part in (company_name, job_title) if part))
        run.font.size = Pt(10)
        doc.add_paragraph()

    for block in cover_letter.strip().split("\n\n"):
        lines = [line.strip() for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        paragraph = doc.add_paragraph("\n".join(lines))
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.15

    filename = f"{uuid.uuid4()}_lettre_motivation.docx"
    output_path = settings.session_output_path(session_id) / filename
    save_docx(doc, str(output_path))

    return filename, f"/api/download/{filename}"


def export_cover_letter_pdf(
    session_id: str,
    cover_letter: str,
    *,
    company_name: str = "",
    job_title: str = "",
) -> tuple[str, str] | None:
    docx_name, _ = export_cover_letter_docx(
        session_id,
        cover_letter,
        company_name=company_name,
        job_title=job_title,
    )
    docx_path = settings.session_output_path(session_id) / docx_name
    pdf_name = docx_name.replace(".docx", ".pdf")
    pdf_path = settings.session_output_path(session_id) / pdf_name
    if export_docx_to_pdf(docx_path, pdf_path):
        return pdf_name, f"/api/download/{pdf_name}"
    return None
