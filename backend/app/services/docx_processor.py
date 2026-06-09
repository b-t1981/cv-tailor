import re
from collections.abc import Iterator
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from app.models.schemas import ParagraphInfo

HEADING_STYLES = {
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "Titre 1",
    "Titre 2",
    "Titre 3",
    "Title",
    "Subtitle",
}

BlockRef = tuple[Paragraph, int | None, str]


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name in HEADING_STYLES or (
        paragraph.runs
        and paragraph.runs[0].bold
        and len(paragraph.text.strip()) < 80
        and not paragraph.text.strip().startswith("•")
        and not paragraph.text.strip().startswith("-")
    )


def _is_heading_text(text: str) -> bool:
    stripped = text.strip()
    return len(stripped) < 80 and stripped.isupper() and len(stripped) > 2


def _expand_lines(prefix: str, text: str) -> list[tuple[str, int | None, str]]:
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return []
    if len(lines) == 1:
        return [(prefix, None, lines[0])]
    return [(f"{prefix}_l{line_idx}", line_idx, line) for line_idx, line in enumerate(lines)]


def _iter_document_blocks(doc: Document) -> Iterator[tuple[str, Paragraph, int | None, str]]:
    for index, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        for block_id, line_idx, line in _expand_lines(f"b{index}", paragraph.text):
            yield block_id, paragraph, line_idx, line

    for table_idx, table in enumerate(doc.tables):
        for cell_idx, tc in enumerate(table._tbl.iter(qn("w:tc"))):
            cell = _Cell(tc, table)
            for para_idx, cell_paragraph in enumerate(cell.paragraphs):
                if not cell_paragraph.text.strip():
                    continue
                prefix = f"t{table_idx}_c{cell_idx}_p{para_idx}"
                for block_id, line_idx, line in _expand_lines(prefix, cell_paragraph.text):
                    yield block_id, cell_paragraph, line_idx, line

    for section_idx, section in enumerate(doc.sections):
        for part_name, part in (("header", section.header), ("footer", section.footer)):
            for para_idx, part_paragraph in enumerate(part.paragraphs):
                if not part_paragraph.text.strip():
                    continue
                prefix = f"h{section_idx}_{part_name}_p{para_idx}"
                for block_id, line_idx, line in _expand_lines(prefix, part_paragraph.text):
                    yield block_id, part_paragraph, line_idx, line


def _build_block_index(doc: Document) -> dict[str, BlockRef]:
    return {
        block_id: (paragraph, line_idx, line_text)
        for block_id, paragraph, line_idx, line_text in _iter_document_blocks(doc)
    }


def extract_paragraphs(doc: Document) -> list[ParagraphInfo]:
    paragraphs: list[ParagraphInfo] = []
    for block_id, paragraph, line_idx, line_text in _iter_document_blocks(doc):
        style_name = paragraph.style.name if paragraph.style else None
        paragraphs.append(
            ParagraphInfo(
                id=block_id,
                text=line_text,
                style=style_name,
                is_heading=_is_heading(paragraph) or _is_heading_text(line_text),
            )
        )
    return paragraphs


def paragraphs_to_prompt_text(paragraphs: list[ParagraphInfo]) -> str:
    lines = []
    for paragraph in paragraphs:
        tag = "[HEADING]" if paragraph.is_heading else "[TEXT]"
        lines.append(f'{paragraph.id} {tag}: "{paragraph.text}"')
    return "\n".join(lines)


def paragraphs_to_readable_cv(paragraphs: list[ParagraphInfo]) -> str:
    """Plain CV text for match scoring (no block IDs)."""
    return "\n".join(paragraph.text for paragraph in paragraphs if paragraph.text.strip())


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> bool:
    if paragraph.text.strip() == new_text.strip():
        return False

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
    return True


def _replace_line_in_paragraph(
    paragraph: Paragraph,
    line_idx: int,
    new_line: str,
    expected_old: str,
) -> bool:
    lines = paragraph.text.split("\n")
    if line_idx >= len(lines):
        return False

    current_line = lines[line_idx].strip()
    if current_line != expected_old.strip():
        return False

    if current_line == new_line.strip():
        return False

    lines[line_idx] = new_line
    return _replace_paragraph_text(paragraph, "\n".join(lines))


def _resolve_block_id(block_id: str, block_index: dict[str, BlockRef]) -> str | None:
    if block_id in block_index:
        return block_id

    legacy = re.fullmatch(r"p_(\d+)", block_id)
    if not legacy:
        return None

    ordered_ids = list(block_index.keys())
    legacy_index = int(legacy.group(1))
    if legacy_index >= len(ordered_ids):
        return None
    return ordered_ids[legacy_index]


def apply_modifications(doc: Document, modifications: dict[str, str]) -> int:
    if not modifications:
        return 0

    block_index = _build_block_index(doc)
    applied = 0

    for raw_id, new_text in modifications.items():
        new_text = new_text.strip()
        if not new_text:
            continue

        block_id = _resolve_block_id(raw_id, block_index)
        if not block_id:
            continue

        paragraph, line_idx, original_text = block_index[block_id]

        if line_idx is None:
            if original_text.strip() == new_text:
                continue
            if _replace_paragraph_text(paragraph, new_text):
                applied += 1
        elif _replace_line_in_paragraph(paragraph, line_idx, new_text, expected_old=original_text):
            applied += 1

    return applied


def load_docx(path: str) -> Document:
    # Load from bytes to avoid python-docx table corruption on repeated file opens.
    with open(path, "rb") as handle:
        return Document(BytesIO(handle.read()))


def save_docx(doc: Document, path: str) -> None:
    doc.save(path)
