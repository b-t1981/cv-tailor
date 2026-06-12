from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.services.docx_pdf_renderer import render_docx_to_pdf
from app.services.pdf_exporter import export_docx_to_pdf


def _build_sample_cv_docx(path: Path) -> None:
    doc = Document()
    title = doc.add_paragraph("JEAN DUPONT")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph("Ingénieur support applicatif")

    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.paragraphs[0].add_run("Compétences").bold = True
    left.add_paragraph("Unix, Windows, SQL")
    right.paragraphs[0].add_run("Expérience").bold = True
    right.add_paragraph("Support N2 — gestion incidents")

    doc.save(path)


def test_render_docx_to_pdf_creates_file(tmp_path: Path):
    docx = tmp_path / "cv.docx"
    pdf = tmp_path / "cv.pdf"
    _build_sample_cv_docx(docx)

    assert render_docx_to_pdf(docx, pdf)
    assert pdf.stat().st_size > 500


def test_export_docx_to_pdf_uses_embedded_path(tmp_path: Path):
    docx = tmp_path / "tailored.docx"
    pdf = tmp_path / "tailored.pdf"
    _build_sample_cv_docx(docx)

    assert export_docx_to_pdf(docx, pdf)
    assert pdf.exists()
