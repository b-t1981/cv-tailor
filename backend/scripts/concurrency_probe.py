#!/usr/bin/env python3
"""Probe multi-user / concurrent behaviour against a running API (no LLM)."""

from __future__ import annotations

import json
import sys
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document

API = "http://127.0.0.1:8001/api"
BACKEND_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND_ROOT))


def make_docx(marker: str) -> bytes:
    doc = Document()
    doc.add_heading(f"CV {marker}", level=1)
    doc.add_paragraph(f"Unique marker for user {marker}: experience and skills.")
    doc.add_paragraph("Second paragraph with more professional background details.")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def new_client() -> httpx.Client:
    return httpx.Client(base_url=API, timeout=30.0)


def upload_preview(client: httpx.Client, marker: str) -> dict:
    content = make_docx(marker)
    response = client.post(
        "/preview",
        files={
            "file": (
                f"{marker.lower()}.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    response.raise_for_status()
    return response.json()


def test_health(workers: int = 25) -> dict:
    def hit(_: int) -> int:
        with new_client() as client:
            return client.get("/health").status_code

    start = time.perf_counter()
    with ThreadPoolExecutor(max_workers=workers) as pool:
        codes = list(pool.map(hit, range(workers)))
    elapsed = time.perf_counter() - start
    return {
        "name": "health_concurrent",
        "ok": all(c == 200 for c in codes),
        "workers": workers,
        "elapsed_s": round(elapsed, 2),
    }


def test_stored_cv_isolated_per_session(markers: list[str]) -> dict:
    """Each httpx.Client keeps its own session cookie — storage must not cross-leak."""
    from app.config import settings
    from app.services.cv_storage_service import cv_storage_service

    clients: dict[str, httpx.Client] = {}
    results: list[tuple[str, str]] = []
    errors: list[str] = []

    try:
        for marker in markers:
            clients[marker] = new_client()

        def worker(marker: str) -> None:
            try:
                data = upload_preview(clients[marker], marker)
                text = data["paragraphs"][0]["text"] if data.get("paragraphs") else ""
                results.append((marker, text))
            except Exception as exc:  # noqa: BLE001
                errors.append(f"{marker}: {exc}")

        threads = [threading.Thread(target=worker, args=(m,)) for m in markers]
        start = time.perf_counter()
        for thread in threads:
            thread.start()
        for thread in threads:
            thread.join(timeout=60)
        elapsed = time.perf_counter() - start

        per_session: dict[str, str | None] = {}
        for marker, client in clients.items():
            sid = client.cookies.get("cv_tailor_sid")
            meta = cv_storage_service.load_metadata(sid) if sid else None
            blob = " ".join(p.text for p in meta["paragraphs"]) if meta else ""
            per_session[marker] = next((m for m in markers if m in blob), None)

        isolated = all(per_session[m] == m for m in markers)
        return {
            "name": "stored_cv_per_session",
            "ok": not errors and isolated,
            "elapsed_s": round(elapsed, 2),
            "errors": errors,
            "storage_owner_per_client": per_session,
            "issue": None if isolated else "Session isolation broken",
        }
    finally:
        for client in clients.values():
            client.close()


def test_download_cross_session_denied() -> dict:
    from app.config import settings

    marker_a = "DOWNLOAD_A"
    marker_b = "DOWNLOAD_B"
    with new_client() as client_a, new_client() as client_b:
        upload_preview(client_a, marker_a)
        # Export path not tested without LLM — create a fake file in A's session dir
        sid_a = client_a.cookies.get("cv_tailor_sid")
        if not sid_a:
            return {"name": "download_isolation", "ok": False, "error": "no session A"}

        out_dir = settings.session_output_path(sid_a)
        fake = out_dir / "probe_secret.docx"
        fake.write_bytes(b"secret")
        try:
            denied = client_b.get("/download/probe_secret.docx").status_code == 404
            allowed = client_a.get("/download/probe_secret.docx").status_code == 200
            return {
                "name": "download_isolation",
                "ok": denied and allowed,
                "owner_can_download": allowed,
                "other_session_blocked": denied,
            }
        finally:
            fake.unlink(missing_ok=True)


def test_prompts_locked() -> dict:
    from app.config import settings
    from app.services.prompt_service import prompt_service

    payload = {
        "system_prompt": "Concurrent probe system prompt " + ("x" * 40),
        "user_prompt": "Concurrent probe user prompt " + ("y" * 40),
    }
    with new_client() as client:
        status = client.put("/prompts", json=payload).status_code

    if settings.allow_prompt_writes:
        prompt_service.reset()
        return {
            "name": "prompts_locked",
            "ok": status == 200,
            "status": status,
            "note": "ALLOW_PROMPT_WRITES=true (dev only)",
        }
    return {
        "name": "prompts_locked",
        "ok": status == 403,
        "status": status,
    }


def main() -> int:
    try:
        with new_client() as client:
            client.get("/health").raise_for_status()
    except Exception as exc:
        print(json.dumps({"error": f"API not reachable at {API}: {exc}"}, indent=2))
        return 1

    reports = [
        test_health(),
        test_stored_cv_isolated_per_session(["USER_ALPHA", "USER_BETA", "USER_GAMMA"]),
        test_download_cross_session_denied(),
        test_prompts_locked(),
    ]
    print(json.dumps(reports, indent=2, ensure_ascii=False))
    return 0 if all(r.get("ok") for r in reports) else 1


if __name__ == "__main__":
    raise SystemExit(main())
